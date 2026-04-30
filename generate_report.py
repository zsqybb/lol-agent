from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

style = doc.styles['Normal']
style.font.name = '宋体'
style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.font.size = Pt(12)
style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

def set_heading1(para):
    para.style = doc.styles['Heading 1']

def add_title(doc, text):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(22)
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return para

def add_h1(doc, text):
    para = doc.add_heading(text, level=1)
    for run in para.runs:
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run.font.color.rgb = RGBColor(0, 51, 102)
    return para

def add_h2(doc, text):
    para = doc.add_heading(text, level=2)
    for run in para.runs:
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run.font.color.rgb = RGBColor(0, 76, 153)
    return para

def add_h3(doc, text):
    para = doc.add_heading(text, level=3)
    for run in para.runs:
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run.font.color.rgb = RGBColor(0, 102, 51)
    return para

def add_p(doc, text, indent=False):
    para = doc.add_paragraph(text)
    para.paragraph_format.first_line_indent = Inches(0.3) if indent else Inches(0)
    return para

def add_bullet(doc, text, level=0):
    para = doc.add_paragraph(style='List Bullet')
    para.add_run(text)
    return para

doc.add_page_break()

add_title(doc, '《英雄联盟》游戏数据智能问答系统')
para = doc.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = para.add_run('综合项目报告')
run.font.size = Pt(16)
run.font.name = '黑体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

doc.add_paragraph()
para = doc.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = para.add_run('—— 基于大语言模型的游戏数据查询与智能问答平台')
run.font.size = Pt(14)
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

doc.add_paragraph()
doc.add_paragraph()

para = doc.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = para.add_run('学    院：计算机与信息学院')
run.font.size = Pt(12)

para = doc.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = para.add_run('专    业：计算机科学与技术')
run.font.size = Pt(12)

para = doc.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = para.add_run('学    号：________________')
run.font.size = Pt(12)

para = doc.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = para.add_run('姓    名：________________')
run.font.size = Pt(12)

para = doc.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = para.add_run('指导教师：________________')
run.font.size = Pt(12)

doc.add_page_break()

add_h1(doc, '第1章 系统概述与需求分析')
add_h2(doc, '1.1 项目背景')
add_p(doc, '《英雄联盟》（League of Legends，简称LOL）是由美国Riot Games公司开发的一款多人在线战术竞技游戏（MOBA），自2009年发布以来已成为全球最具影响力的电子竞技项目之一。游戏拥有超过160位英雄角色、丰富的装备系统、复杂的符文天赋机制以及庞大的世界观故事体系。', indent=True)
add_p(doc, '然而，玩家在游戏过程中常常面临以下问题：', indent=True)
add_bullet(doc, '英雄攻略获取困难：玩家需要查询特定英雄的出装、符文、技能加点等攻略，往往需要在多个网站间切换')
add_bullet(doc, '版本信息滞后：网络上的攻略文章可能基于旧版本数据，与当前游戏版本脱节')
add_bullet(doc, '知识分散：英雄背景故事、阵营关系、世界观等 lore 内容分散在多个平台')
add_bullet(doc, '数据分析不足：玩家想知道某个英雄的克制关系、配合英雄等深度信息')
add_p(doc, '本项目旨在构建一个基于大语言模型的《英雄联盟》游戏数据智能问答系统，整合游戏数据、知识库和AI能力，为玩家提供一站式的智能问答服务。', indent=True)

add_h2(doc, '1.2 系统目标与功能定位')
add_p(doc, '本系统定位为游戏玩家的智能助手，而非普通的聊天机器人。系统需要完成以下核心目标：', indent=True)
add_bullet(doc, '游戏数据查询：整合英雄信息、装备数据、符文配置等结构化数据')
add_bullet(doc, '智能问答服务：基于知识库和AI能力，提供准确的英雄攻略、游戏知识回答')
add_bullet(doc, '多模式问答：支持AI自由问答、英雄攻略问答、故事lore问答三种模式')
add_bullet(doc, '对话记忆功能：能够记忆对话上下文，实现多轮连贯对话')
add_bullet(doc, '实时数据更新：后端可扩展接入最新游戏版本数据')

add_h2(doc, '1.3 目标用户分析')
table = doc.add_table(rows=5, cols=2)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
cells_header = table.rows[0].cells
cells_header[0].text = '用户类型'
cells_header[1].text = '使用场景'
for run in cells_header[0].paragraphs[0].runs:
    run.bold = True
for run in cells_header[1].paragraphs[0].runs:
    run.bold = True

data = [
    ('新手玩家', '查询英雄基础信息、推荐出装、快速上手指南'),
    ('进阶玩家', '了解英雄技巧、高阶玩法、版本强势英雄'),
    ('深度玩家', '研究英雄克制关系、阵容搭配、战术分析'),
    (' Lore爱好者', '探索英雄背景故事、阵营历史、世界观设定'),
]
for i, (u, s) in enumerate(data):
    table.rows[i+1].cells[0].text = u
    table.rows[i+1].cells[1].text = s

doc.add_paragraph()

add_h2(doc, '1.4 系统特点')
add_p(doc, '本系统区别于普通大模型Chatbot的关键特点：', indent=True)
add_bullet(doc, '专业领域知识库：构建了包含英雄数据、装备信息、符文配置、Lore故事等内容的专业知识库')
add_bullet(doc, '多技能调度系统：实现了基于关键词和上下文的智能技能调度机制')
add_bullet(doc, '结构化数据处理：后端集成Riot API获取真实玩家数据，提供个性化分析')
add_bullet(doc, '多轮对话记忆：实现了基于文件的会话记忆系统，支持长对话上下文')
add_bullet(doc, '实时内容生成：AI生成的回答结合知识库数据，确保准确性和专业性')

doc.add_page_break()
add_h1(doc, '第2章 系统架构设计')
add_h2(doc, '2.1 整体架构')
add_p(doc, '本系统采用前后端分离的B/S架构，前端负责用户交互和界面展示，后端负责业务逻辑和数据处理。整体架构分为以下层次：', indent=True)
add_bullet(doc, '前端展示层：基于HTML5、CSS3、JavaScript构建的单页应用（SPA）')
add_bullet(doc, '后端服务层：基于Python Flask框架的RESTful API服务')
add_bullet(doc, '技能调度层：基于Skill Dispatcher的智能任务分发系统')
add_bullet(doc, '知识库层：JSON格式的结构化知识库 + 外部API数据源')
add_bullet(doc, '大模型层：集成讯飞星火大模型API提供智能问答能力')

add_h2(doc, '2.2 技术架构图')
add_p(doc, '系统技术架构如下所示：', indent=True)
add_p(doc, '┌─────────────────────────────────────────────────────────────┐', indent=True)
add_p(doc, '│                        前端展示层                          │', indent=True)
add_p(doc, '│   HTML5 + CSS3 + JavaScript + FontAwesome + 响应式设计    │', indent=True)
add_p(doc, '├─────────────────────────────────────────────────────────────┤', indent=True)
add_p(doc, '│                       Flask Web 服务                        │', indent=True)
add_p(doc, '│            /api/player  /api/ai-chat  /api/champions        │', indent=True)
add_p(doc, '├─────────────────────────────────────────────────────────────┤', indent=True)
add_p(doc, '│                      Skill 调度层                          │', indent=True)
add_p(doc, '│     GeneralSkill  GameplaySkill  LoreSkill  PlayerSkill    │', indent=True)
add_p(doc, '├───────────────────┬───────────────────────────────────────┤', indent=True)
add_p(doc, '│    知识库层        │              数据源层                 │', indent=True)
add_p(doc, '│  knowledge_base    │   Riot API   讯飞星火API   本地数据   │', indent=True)
add_p(doc, '│  lore/*.json       │                                       │', indent=True)
add_p(doc, '└───────────────────┴───────────────────────────────────────┘', indent=True)

add_h2(doc, '2.3 核心模块设计')
add_h3(doc, '2.3.1 前端模块')
add_p(doc, '前端采用模块化JavaScript架构，主要模块包括：', indent=True)
add_bullet(doc, 'app.js：核心应用逻辑，包括页面切换、数据渲染、API调用')
add_bullet(doc, 'UI组件：英雄卡片、属性条、出装列表、符文展示等')
add_bullet(doc, '聊天系统：三合一聊天界面，支持AI/英雄/故事三种模式切换')
add_bullet(doc, '状态管理：会话状态、历史记录、加载状态管理')

add_h3(doc, '2.3.2 后端模块')
add_p(doc, '后端采用Flask框架，核心路由和功能模块：', indent=True)
add_bullet(doc, 'web_server.py：主服务器，处理所有API请求')
add_bullet(doc, 'riot_api_client.py：Riot API客户端，获取玩家数据')
add_bullet(doc, 'chat_memory.py：对话记忆管理，实现会话状态持久化')
add_bullet(doc, 'skills/：技能模块目录，包含各类专业技能封装')

add_h3(doc, '2.3.3 数据文件结构')
add_p(doc, '项目目录结构如下：', indent=True)
add_p(doc, 'f:\\code-tengxun\\', indent=True)
add_p(doc, '├── index.html          # 网站主页面', indent=True)
add_p(doc, '├── style.css          # 样式表', indent=True)
add_p(doc, '├── app.js             # 前端JavaScript', indent=True)
add_p(doc, '├── web_server.py      # Flask后端服务器', indent=True)
add_p(doc, '├── riot_api_client.py # Riot API客户端', indent=True)
add_p(doc, '├── chat_memory.py     # 对话记忆管理', indent=True)
add_p(doc, '├── knowledge_base.json # 游戏知识库', indent=True)
add_p(doc, '├── skills/            # 技能模块', indent=True)
add_p(doc, '│   ├── skill_dispatcher.py', indent=True)
add_p(doc, '│   ├── general_skill.py', indent=True)
add_p(doc, '│   ├── gameplay_skill.py', indent=True)
add_p(doc, '│   ├── lore_skill.py', indent=True)
add_p(doc, '│   └── player_skill.py', indent=True)
add_p(doc, '└── lore/              # 背景故事数据', indent=True)
add_p(doc, '    └── factions/*.json', indent=True)

add_h2(doc, '2.4 业务流程设计')
add_p(doc, '用户请求处理的完整业务流程：', indent=True)
add_bullet(doc, '用户在前端界面输入问题或选择功能')
add_bullet(doc, '前端根据功能类型调用相应API（如/api/ai-chat、/api/player）')
add_bullet(doc, '后端接收请求，根据参数判断是否需要调度Skill')
add_bullet(doc, 'Skill Dispatcher分析问题关键词和上下文，选择合适的Skill处理')
add_bullet(doc, 'Skill从知识库或外部API获取数据，进行处理和格式化')
add_bullet(doc, '如果需要AI生成，结合历史记忆调用大模型API')
add_bullet(doc, '后端返回结构化JSON响应给前端')
add_bullet(doc, '前端渲染数据，更新界面显示')

doc.add_page_break()
add_h1(doc, '第3章 NLP技术应用与实现')
add_h2(doc, '3.1 主要NLP任务识别')
add_p(doc, '本系统涉及以下NLP典型任务：', indent=True)

table = doc.add_table(rows=7, cols=3)
table.style = 'Table Grid'
headers = ['NLP任务', '具体应用', '实现方式']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
    for run in table.rows[0].cells[i].paragraphs[0].runs:
        run.bold = True

nlp_data = [
    ('意图识别', '判断用户问题是查询英雄、查询玩家、还是闲聊', '关键词匹配 + 上下文分析'),
    ('命名实体识别', '从问题中提取英雄名、装备名、地名等实体', '模糊匹配 + 知识库映射'),
    ('语义检索', '在知识库中查找相关内容', '关键词检索 + 语义相似度'),
    ('文本生成', '生成结构化的攻略回答', '大模型 + 模板填充'),
    ('情感分析', '分析玩家评论情感倾向', '关键词 + 规则模式'),
    ('对话管理', '维护多轮对话上下文', '会话记忆 + 状态跟踪'),
]
for i, (task, app, method) in enumerate(nlp_data):
    table.rows[i+1].cells[0].text = task
    table.rows[i+1].cells[1].text = app
    table.rows[i+1].cells[2].text = method

doc.add_paragraph()

add_h2(doc, '3.2 意图识别与分类')
add_p(doc, '系统实现了基于规则的意图识别机制，通过多级匹配策略判断用户问题类型：', indent=True)
add_bullet(doc, '一级匹配：检测强制技能调度标识（如"@英雄"、英雄名模式）')
add_bullet(doc, '二级匹配：检测关键词集合（如"出装"触发GameplaySkill）')
add_bullet(doc, '三级匹配：基于模糊相似度匹配英雄名、装备名等实体')
add_bullet(doc, '兜底策略：使用GeneralSkill进行通用问答')

add_p(doc, 'Intent Classification流程：', indent=True)
add_p(doc, '用户输入 → 预处理（分词、去噪）→ 关键词检测 → 模式匹配 → 技能调度', indent=True)

add_h2(doc, '3.3 命名实体识别')
add_p(doc, '系统通过以下方式实现NER：', indent=True)
add_bullet(doc, '英雄名实体：从知识库获取所有英雄名称，支持中英文别名')
add_bullet(doc, '装备名实体：从知识库获取装备列表')
add_bullet(doc, '阵营名实体：从lore知识库获取所有阵营名称')
add_p(doc, '实体匹配采用fuzz.ratio进行模糊匹配，阈值为80%，确保良好的容错性。', indent=True)

add_h2(doc, '3.4 语义检索与知识库问答')
add_p(doc, '当用户询问英雄攻略或游戏知识时，系统从知识库中检索相关内容：', indent=True)
add_bullet(doc, '关键词索引：知识库按类别（heroes、items、runes等）组织')
add_bullet(doc, '相似度匹配：使用fuzzywuzzy库进行字符串相似度计算')
add_bullet(doc, '上下文扩展：当检测到上下文指代时（如"他"指代上一轮提到的英雄），利用last_champ_id扩展检索范围')

add_h2(doc, '3.5 大模型文本生成')
add_p(doc, '系统集成讯飞星火大模型进行文本生成，主要应用场景：', indent=True)
add_bullet(doc, '通用问答：当问题不属于特定知识领域时，由GeneralSkill调用大模型回答')
add_bullet(doc, '攻略生成：结合知识库数据，大模型生成结构化的英雄攻略')
add_bullet(doc, '对话续写：利用对话历史记忆，大模型生成连贯的对话回复')
add_bullet(doc, '记忆压缩：当对话超过5轮时，大模型负责总结压缩历史记忆')

doc.add_page_break()
add_h1(doc, '第4章 Agent调度与Skills封装')
add_h2(doc, '4.1 Agent调度机制设计')
add_p(doc, '本系统采用基于Skill Dispatcher的Agent调度机制，实现任务的智能分发和处理。调度器维护一个技能注册表，根据预定义的规则将用户请求分发到对应的技能模块。', indent=True)

add_h2(doc, '4.2 Skill调度器实现')
add_p(doc, 'Skill Dispatcher的核心调度逻辑：', indent=True)
add_p(doc, '┌────────────────────────────────────────────────────┐', indent=True)
add_p(doc, '│              Skill Dispatcher                     │', indent=True)
add_p(doc, '├────────────────────────────────────────────────┤', indent=True)
add_p(doc, '│  1. 接收用户消息和上下文                          │', indent=True)
add_p(doc, '│  2. 检查强制调度标识（dispatch_skill参数）        │', indent=True)
add_p(doc, '│  3. 分析消息，提取关键词                          │', indent=True)
add_p(doc, '│  4. 遍历Skills，按优先级匹配                      │', indent=True)
add_p(doc, '│  5. 调用匹配的Skill处理请求                       │', indent=True)
add_p(doc, '│  6. 返回处理结果                                  │', indent=True)
add_p(doc, '└────────────────────────────────────────────────┘', indent=True)

add_h2(doc, '4.3 Skills模块设计')
add_p(doc, '系统实现了四个核心Skill模块：', indent=True)

add_h3(doc, '4.3.1 GeneralSkill（通用技能）')
add_bullet(doc, '职责：处理通用问答、不属于特定领域的问题')
add_bullet(doc, '优先级：最低（兜底技能）')
add_bullet(doc, '调用大模型：讯飞星火API')
add_bullet(doc, '提示词优化：包含角色定位、回答风格、知识边界设定')

add_h3(doc, '4.3.2 GameplaySkill（游戏攻略技能）')
add_bullet(doc, '职责：处理英雄出装、符文、技能加点、技巧等攻略问题')
add_bullet(doc, '触发关键词：出装、符文、技巧、技能、加点、对线、打野')
add_bullet(doc, '数据来源：knowledge_base.json英雄数据')
add_bullet(doc, '输出格式：结构化攻略，包含装备推荐、符文配置、技巧要点')

add_h3(doc, '4.3.3 LoreSkill（背景故事技能）')
add_bullet(doc, '职责：处理英雄背景故事、阵营介绍、世界观问题')
add_bullet(doc, '触发关键词：背景、故事、历史、来历、阵营、传说')
add_bullet(doc, '数据来源：lore/factions/*.json阵营故事数据')
add_bullet(doc, '输出格式：故事性描述，包含背景、关系、关键事件')

add_h3(doc, '4.3.4 PlayerSkill（玩家数据技能）')
add_bullet(doc, '职责：处理玩家数据查询，如战绩、排名、KDA等')
add_bullet(doc, '触发关键词：战绩、段位、排名、比赛、击杀')
add_bullet(doc, '数据来源：Riot API真实玩家数据')
add_bullet(doc, '输出格式：玩家数据卡片，包含统计数据和趋势分析')

add_h2(doc, '4.4 强制调度机制')
add_p(doc, '系统支持强制调度，即在API调用时直接指定使用某个Skill：', indent=True)
add_p(doc, 'mode参数：ai（默认通用）、hero（强制英雄攻略）、lore（强制故事问答）', indent=True)
add_p(doc, 'dispatch_skill参数：直接指定技能名称进行调度', indent=True)
add_p(doc, 'champ_id参数：传递当前英雄上下文，用于解决指代问题', indent=True)

add_h2(doc, '4.5 Skills封装示例')
add_p(doc, '以GameplaySkill为例，其结构化的回答格式如下：', indent=True)
add_p(doc, '📍 定位：战士 / 刺客', indent=True)
add_p(doc, '📊 属性：攻击 | 防御 | 魔法 | 难度', indent=True)
add_p(doc, '', indent=True)
add_p(doc, '🎯 **推荐出装**', indent=True)
add_p(doc, '1. 核心装备：xxx', indent=True)
add_p(doc, '2. 鞋子：xxx', indent=True)
add_p(doc, '', indent=True)
add_p(doc, '🎴 **推荐符文**', indent=True)
add_p(doc, '主系：精密 - xxx', indent=True)
add_p(doc, '副系：坚决 - xxx', indent=True)
add_p(doc, '', indent=True)
add_p(doc, '💡 **技巧要点**', indent=True)
add_p(doc, '1. xxx', indent=True)

doc.add_page_break()
add_h1(doc, '第5章 知识库构建与技术实现')
add_h2(doc, '5.1 知识库设计')
add_p(doc, '本系统的知识库包含两个主要部分：结构化知识库和Lore故事库。', indent=True)

add_h2(doc, '5.1.1 结构化知识库（knowledge_base.json）')
add_p(doc, 'knowledge_base.json包含以下内容：', indent=True)
add_bullet(doc, 'heroes：英雄基础数据（名称、称号、定位、属性、技能）')
add_bullet(doc, 'items：装备数据（名称、价格、属性、合成路径）')
add_bullet(doc, 'runes：符文数据（名称、效果、组合）')
add_bullet(doc, 'matchups：克制关系数据（英雄间的克制/被克制关系）')
add_bullet(doc, 'tips：游戏技巧问答对')

add_h2(doc, '5.1.2 Lore故事库（lore/）')
add_p(doc, 'lore目录包含各阵营的背景故事JSON文件：', indent=True)
add_bullet(doc, 'demacia.json：德玛西亚 - 崇尚荣耀与正义的王国')
add_bullet(doc, 'noxus.json：诺克萨斯 - 扩张主义帝国')
add_bullet(doc, 'ionia.json：艾欧尼亚 - 追求平衡的群岛')
add_bullet(doc, 'freljord.json：弗雷尔卓德 - 冰霜之地')
add_bullet(doc, 'zaun.json：祖安 - 地下工业城邦')
add_bullet(doc, 'piltover.json：皮尔特沃夫 - 进步之城')
add_bullet(doc, 'bandle-city.json：班德尔城 - 约德尔人故乡')
add_bullet(doc, 'shadow-isles.json：暗影岛 - 被诅咒的黑雾岛屿')
add_bullet(doc, 'ixtal.json：以绪塔尔 - 元素魔法国度')

add_h2(doc, '5.2 Riot API集成')
add_p(doc, '系统通过riot_api_client.py模块集成Riot API，获取真实玩家数据：', indent=True)
add_bullet(doc, '账户信息：获取玩家游戏名、等级、头像')
add_bullet(doc, '排位数据：获取段位、胜点、排名')
add_bullet(doc, '比赛记录：获取最近比赛详情、英雄、KDA、出装')
add_p(doc, 'API支持多个端点：Asia、Americas、Europe等区域。', indent=True)

add_h2(doc, '5.3 对话记忆系统实现')
add_p(doc, '系统实现了基于文件的会话记忆管理（chat_memory.py）：', indent=True)
add_bullet(doc, '会话创建：新对话创建唯一的session_id，生成对应内存文件')
add_bullet(doc, '消息存储：每轮对话追加写入内存文件')
add_bullet(doc, '记忆压缩：超过5轮对话时，调用大模型总结压缩历史')
add_bullet(doc, '会话清理：对话结束时删除内存文件')

add_p(doc, 'Memory压缩示例：', indent=True)
add_p(doc, '压缩前：用户问"亚索怎么出装" → AI回答出装推荐', indent=True)
add_p(doc, '压缩后：用户问"他适合什么符文" → AI知道"他"指亚索', indent=True)

add_h2(doc, '5.4 前端交互实现')
add_p(doc, '前端采用现代化的单页应用架构：', indent=True)
add_bullet(doc, 'Tab切换：三合一聊天界面通过Tab切换AI/英雄/故事模式')
add_bullet(doc, '实时反馈：打字动画、加载状态、错误提示')
add_bullet(doc, '响应式设计：支持桌面和移动端访问')
add_bullet(doc, '无刷新交互：局部更新DOM，不重新加载页面')

add_h2(doc, '5.5 后端API设计')
add_p(doc, '系统主要API端点：', indent=True)
table = doc.add_table(rows=6, cols=3)
table.style = 'Table Grid'
for i, h in enumerate(['端点', '方法', '功能']):
    table.rows[0].cells[i].text = h
    for run in table.rows[0].cells[i].paragraphs[0].runs:
        run.bold = True

api_data = [
    ('/api/ai-chat', 'POST', '智能问答API，支持mode参数指定模式'),
    ('/api/player', 'GET', '查询玩家数据，需要game_name和tag_line参数'),
    ('/api/champions', 'GET', '获取英雄列表'),
    ('/api/champion/<id>', 'GET', '获取英雄详情'),
    ('/api/chat/session', 'POST', '创建新对话会话'),
]
for i, (ep, method, func) in enumerate(api_data):
    table.rows[i+1].cells[0].text = ep
    table.rows[i+1].cells[1].text = method
    table.rows[i+1].cells[2].text = func

doc.add_paragraph()

doc.add_page_break()
add_h1(doc, '第6章 系统特色与价值分析')
add_h2(doc, '6.1 系统特色总结')
add_p(doc, '本系统相比普通大模型Chatbot具有以下特色：', indent=True)

add_h3(doc, '6.1.1 专业领域深度整合')
add_bullet(doc, '构建了完整的《英雄联盟》领域知识库，涵盖160+英雄数据')
add_bullet(doc, '整合了游戏世界观故事，支持Lore问答')
add_bullet(doc, '接入Riot API，获取真实玩家数据')

add_h3(doc, '6.1.2 智能任务调度')
add_bullet(doc, '实现了基于优先级的多Skill调度机制')
add_bullet(doc, '支持强制调度和上下文感知')
add_bullet(doc, '各Skill职责明确，输出格式统一')

add_h3(doc, '6.1.3 多轮对话记忆')
add_bullet(doc, '基于文件的会话状态管理')
add_bullet(doc, '5轮对话自动压缩，保持上下文简洁')
add_bullet(doc, '支持指代消解（如"他"指代上一轮提到的英雄）')

add_h3(doc, '6.1.4 现代化用户体验')
add_bullet(doc, '三合一聊天界面，模式切换便捷')
add_bullet(doc, '响应式设计，支持多设备访问')
add_bullet(doc, '实时打字动画，操作反馈友好')

add_h2(doc, '6.2 技术创新点')
add_bullet(doc, '多模式问答融合：将通用问答、英雄攻略、故事问答三种模式整合到统一界面')
add_bullet(doc, 'Skill级联调度：支持强制调度与智能匹配结合的调度策略')
add_bullet(doc, '记忆压缩机制：利用大模型实现对话历史的自动总结与压缩')
add_bullet(doc, '领域知识增强：通过知识库检索增强大模型的专业知识输出')

add_h2(doc, '6.3 应用价值')
add_p(doc, '本系统具有以下应用价值：', indent=True)
add_bullet(doc, '玩家助手：为《英雄联盟》玩家提供便捷的游戏知识查询服务')
add_bullet(doc, '学习工具：帮助新手玩家快速了解游戏机制和英雄攻略')
add_bullet(doc, '社区应用：可扩展为游戏论坛、直播平台的智能客服')
add_bullet(doc, '技术示范：展示了大模型在游戏领域应用的技术路径')

add_h2(doc, '6.4 与普通Chatbot对比')
table = doc.add_table(rows=6, cols=3)
table.style = 'Table Grid'
for i, h in enumerate(['对比维度', '普通Chatbot', '本系统']):
    table.rows[0].cells[i].text = h
    for run in table.rows[0].cells[i].paragraphs[0].runs:
        run.bold = True

compare_data = [
    ('知识来源', '训练数据（可能过时）', '实时知识库 + Riot API'),
    ('领域专业性', '泛化知识', '英雄联盟专业领域知识'),
    ('数据准确性', '可能幻觉', '知识库检索 + 结构化输出'),
    ('交互方式', '单一对话', '多模式（AI/英雄/故事）'),
    ('上下文记忆', '有限上下文窗口', '长期记忆 + 自动压缩'),
]
for i, (dim, chatbot, our) in enumerate(compare_data):
    table.rows[i+1].cells[0].text = dim
    table.rows[i+1].cells[1].text = chatbot
    table.rows[i+1].cells[2].text = our

doc.add_paragraph()

add_h2(doc, '6.5 改进方向')
add_p(doc, '未来可从以下方向进行改进：', indent=True)
add_bullet(doc, '多模态扩展：加入英雄皮肤图片、比赛视频片段等视觉内容')
add_bullet(doc, '实时数据接入：接入Riot API获取最新版本数据')
add_bullet(doc, '个性化推荐：基于用户游戏风格推荐英雄和出装')
add_bullet(doc, '社区集成：接入Reddit、Twitter等社区数据，分析玩家舆情')
add_bullet(doc, '语音交互：加入语音识别和语音播报功能')

doc.add_page_break()
add_h1(doc, '结论')
add_p(doc, '本项目成功实现了一个基于大语言模型的《英雄联盟》游戏数据智能问答系统。系统综合运用了大语言模型、RAG知识库问答、Agent任务调度、Skills能力封装等自然语言处理技术，实现了区别于普通Chatbot的专业领域智能问答系统。', indent=True)
add_p(doc, '通过本次综合实践，我们深入学习了自然语言处理和大模型应用开发的相关技术，包括：', indent=True)
add_bullet(doc, '如何设计并实现多Skill调度机制')
add_bullet(doc, '如何构建和管理领域知识库')
add_bullet(doc, '如何实现多轮对话的记忆和管理')
add_bullet(doc, '如何将大模型能力与专业领域知识结合')
add_bullet(doc, '如何构建现代化的Web应用界面')
add_p(doc, '本系统已能够实际运行并提供智能问答服务，在英雄联盟游戏知识问答领域具有一定的应用价值和参考意义。', indent=True)

doc.add_page_break()
add_h1(doc, '参考文献')
add_p(doc, '[1] Riot Games. League of Legends Developer API Documentation. https://developer.riotgames.com/', indent=True)
add_p(doc, '[2] 讯飞星火大模型API文档. https://spark-api.xf-yun.com/', indent=True)
add_p(doc, '[3] Flask Web框架文档. https://flask.palletsprojects.com/', indent=True)
add_p(doc, '[4] Riot Games. League of Legends Universe - Lore & Champions. https://universe.leagueoflegends.com/', indent=True)
add_p(doc, '[5] Bilibili. 自然语言处理综合实践课程资料. https://mp.weixin.qq.com/s/4sWzc6v6I2f3DE7jkTLBhQ', indent=True)

doc.save(r'f:\code-tengxun\LOL智能问答系统_综合项目报告.docx')
print('报告已生成: f:\\code-tengxun\\LOL智能问答系统_综合项目报告.docx')
